"""
Generador de documents Word (apunts d'estudi i guia de presentació).
Crea el fitxer .docx amb contingut complet per estudiar i exposar.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

from processors.content_processor import PresentationPlan


def create_study_guide(plan: PresentationPlan, output_path: str | Path) -> Path:
    """
    Crea un document Word amb apunts ultra-comprimits (màx 5 pàgines).

    Args:
        plan: Pla de presentació amb totes les diapositives.
        output_path: Ruta on guardar el fitxer .docx.

    Returns:
        Path al fitxer creat.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Configurar estils compactes
    _setup_styles(doc)

    # ═══════════════════════════════════════════════════════════════════
    # PORTADA ULTRA-COMPACTA
    # ═══════════════════════════════════════════════════════════════════
    title = doc.add_heading(f"APUNTS: {plan.chapter_title}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Grup {plan.group_name} | {plan.chapter_name} | {len(plan.slides)} slides")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ═══════════════════════════════════════════════════════════════════
    # APUNTS MEGA-SINTETITZATS (1-2 pàgines)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("APUNTS SINTETITZATS", level=1)

    # Resum ultra-compact
    if plan.study_summary:
        summary_p = doc.add_paragraph()
        summary_p.add_run("RESUM: ").font.bold = True
        summary_p.add_run(plan.study_summary[:500] + "..." if len(plan.study_summary) > 500 else plan.study_summary)
        summary_p.paragraph_format.space_after = Pt(6)

    # Contingut sintetitzat per seccions
    current_section = ""
    section_content = []

    for slide in plan.slides:
        if slide.slide_type in ["title", "index"]:
            continue

        # Agrupar per secció
        section_title = slide.title.split(":")[0].split(".")[0] if ":" in slide.title or "." in slide.title else slide.title[:30]

        if section_title != current_section:
            # Imprimir secció anterior
            if section_content:
                _add_compact_section(doc, current_section, section_content)
            current_section = section_title
            section_content = []

        # Afegir contingut sintetitzat
        compact_content = []
        for item in slide.content:
            if item and item.strip():
                # Simplificar i comprimir
                text = item.strip().replace("**", "").replace("- ", "").replace("• ", "")
                if text and len(text) > 3:  # Evitar línies massa curtes
                    compact_content.append(text[:100])  # Limitar longitud

        if compact_content:
            section_content.extend(compact_content[:3])  # Màxim 3 punts per slide

    # Última secció
    if section_content:
        _add_compact_section(doc, current_section, section_content)

    # Conceptes clau ultra-comprimits
    if plan.key_concepts:
        doc.add_heading("CONCEPTES CLAU", level=2)
        for concept in plan.key_concepts[:10]:  # Màxim 10 conceptes
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(concept[:80]).font.size = Pt(10)  # Compact
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════
    # GUIA DE PRESENTACIÓ ULTRA-COMPACTA (2-3 pàgines)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("GUIA DE PRESENTACIÓ", level=1)

    # Temps total
    total_seconds = sum(s.duration_seconds for s in plan.slides)
    total_min = total_seconds // 60
    p = doc.add_paragraph()
    p.add_run(f"TEMPS TOTAL: {total_min} min ({len(plan.slides)} slides)").font.bold = True
    p.paragraph_format.space_after = Pt(6)

    for slide in plan.slides:
        # Format ultra-compact per slide
        slide_header = doc.add_paragraph()
        slide_header.paragraph_format.space_after = Pt(3)

        # Número i títol en una línia
        num_run = slide_header.add_run(f"#{slide.number} ")
        num_run.font.bold = True
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = RGBColor(224, 122, 47)

        title_run = slide_header.add_run(slide.title)
        title_run.font.bold = True
        title_run.font.size = Pt(11)

        # Temps
        time_min = slide.duration_seconds // 60
        time_str = f" ({time_min}min)" if time_min else f" ({slide.duration_seconds}s)"
        time_run = slide_header.add_run(time_str)
        time_run.font.size = Pt(9)
        time_run.font.color.rgb = RGBColor(128, 128, 128)

        # Punts clau sintetitzats (màxim 3)
        key_points = []
        for item in slide.content[:3]:  # Només primers 3 punts
            if item and item.strip():
                text = item.strip().replace("**", "").replace("- ", "").replace("• ", "")
                if text and len(text) > 2:
                    key_points.append(text[:60])  # Ultra-compact

        if key_points:
            points_p = doc.add_paragraph()
            points_p.paragraph_format.left_indent = Inches(0.2)
            points_p.paragraph_format.space_after = Pt(2)
            for point in key_points:
                points_p.add_run("• " + point + " ").font.size = Pt(9)

        # Notes ultra-sintetitzades
        if slide.speaker_notes:
            notes_p = doc.add_paragraph()
            notes_p.paragraph_format.left_indent = Inches(0.2)
            notes_p.paragraph_format.space_after = Pt(2)
            # Comprimir notes a 1-2 frases
            notes = slide.speaker_notes
            if len(notes) > 200:
                notes = notes[:200] + "..."
            notes_run = notes_p.add_run(notes)
            notes_run.font.size = Pt(9)
            notes_run.font.italic = True
            notes_run.font.color.rgb = RGBColor(64, 64, 64)

        # Separador mínim
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(6)

    # Guardar
    doc.save(str(output_path))
    print(f"Xuleta guardada: {output_path}")
    return output_path


def _add_compact_section(doc, section_title, content_list):
    """Afegeix una secció compacta amb bullet points."""
    if not content_list:
        return

    # Títol de secció
    section_p = doc.add_paragraph()
    section_run = section_p.add_run(section_title)
    section_run.font.bold = True
    section_run.font.size = Pt(12)
    section_run.font.color.rgb = RGBColor(224, 122, 47)
    section_p.paragraph_format.space_after = Pt(3)

    # Contingut com bullet points compactes
    for item in content_list[:8]:  # Màxim 8 punts per secció
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item).font.size = Pt(10)

    # Espai mínim
    doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # PART 1: APUNTS COMPLETS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("PART 1: APUNTS DEL TEMA", level=1)

    intro = doc.add_paragraph()
    intro.add_run("Aquests apunts expliquen tots els conceptes del capítol de forma detallada. "
                  "Llegeix-los per entendre el tema abans de preparar la presentació.").font.size = Pt(11)
    doc.add_paragraph()

    # Resum inicial
    if plan.study_summary:
        doc.add_heading("Introducció al tema", level=2)
        summary = doc.add_paragraph(plan.study_summary)
        summary.style = 'Normal'
        doc.add_paragraph()

    # Contingut de cada diapositiva com a apunts
    current_section = ""
    for slide in plan.slides:
        # Saltar portada i índex
        if slide.slide_type in ["title", "index"]:
            continue

        # Nou apartat si el títol canvia significativament
        section_title = slide.title.split(":")[0] if ":" in slide.title else slide.title

        # Afegir títol de secció
        doc.add_heading(slide.title, level=2)

        # Contingut explicatiu complet
        if slide.content:
            for item in slide.content:
                if not item or not item.strip():
                    continue

                # Processar el contingut
                text = item.strip()

                # Detectar si és un subtítol (majúscules)
                if text.isupper() and len(text) < 50:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(224, 122, 47)
                    continue

                # Detectar si és una llista numerada
                if text and text[0].isdigit() and len(text) > 1 and text[1] == '.':
                    p = doc.add_paragraph(text, style='List Number')
                    p.paragraph_format.left_indent = Inches(0.3)
                    continue

                # Detectar si és un bullet point
                if text.startswith("-") or text.startswith("•"):
                    clean_text = text.lstrip("-•").strip()
                    p = doc.add_paragraph(clean_text, style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.3)
                    continue

                # Text normal - processar negreta
                p = doc.add_paragraph()
                _add_text_with_bold(p, text)

        # Explicació addicional del speaker_notes
        if slide.speaker_notes:
            p = doc.add_paragraph()
            run = p.add_run("Explicació: ")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(139, 90, 43)

            explanation = p.add_run(slide.speaker_notes)
            explanation.font.size = Pt(11)
            explanation.font.color.rgb = RGBColor(64, 64, 64)

        doc.add_paragraph()  # Espai entre seccions

    # Conceptes clau
    if plan.key_concepts:
        doc.add_heading("Conceptes clau per recordar", level=2)
        for concept in plan.key_concepts:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(concept)
            run.font.size = Pt(11)
            run.font.bold = True
        doc.add_paragraph()

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # PART 2: GUIA DE PRESENTACIÓ
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("PART 2: GUIA DE PRESENTACIÓ", level=1)

    intro2 = doc.add_paragraph()
    intro2.add_run("Aquesta guia t'indica què dir a cada diapositiva i com exposar-ho. "
                   "Utilitza-la per practicar i durant la presentació.").font.size = Pt(11)
    doc.add_paragraph()

    # Temps total
    total_seconds = sum(s.duration_seconds for s in plan.slides)
    total_min = total_seconds // 60
    p = doc.add_paragraph()
    run = p.add_run(f"Temps total estimat: {total_min} minuts ({len(plan.slides)} diapositives)")
    run.font.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()

    for slide in plan.slides:
        # Capçalera de la slide
        slide_heading = doc.add_paragraph()

        # Número amb fons destacat
        number_run = slide_heading.add_run(f"SLIDE {slide.number}")
        number_run.font.bold = True
        number_run.font.color.rgb = RGBColor(255, 255, 255)
        number_run.font.size = Pt(11)
        # No podem fer fons, així que fem servir color taronja
        number_run.font.color.rgb = RGBColor(224, 122, 47)

        slide_heading.add_run(" | ")

        title_run = slide_heading.add_run(slide.title)
        title_run.font.bold = True
        title_run.font.size = Pt(12)

        # Temps
        time_min = slide.duration_seconds // 60
        time_sec = slide.duration_seconds % 60
        if time_min:
            time_str = f" ({time_min} min {time_sec}s)" if time_sec else f" ({time_min} min)"
        else:
            time_str = f" ({time_sec}s)"
        time_run = slide_heading.add_run(time_str)
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = RGBColor(128, 128, 128)

        # Tipus de slide
        type_map = {
            "title": "Portada",
            "index": "Índex",
            "content": "Contingut",
            "conclusion": "Conclusió"
        }
        slide_type_name = type_map.get(slide.slide_type, slide.slide_type)

        type_p = doc.add_paragraph()
        type_p.paragraph_format.left_indent = Inches(0.2)
        type_run = type_p.add_run(f"Tipus: {slide_type_name}")
        type_run.font.size = Pt(10)
        type_run.font.italic = True
        type_run.font.color.rgb = RGBColor(128, 128, 128)

        # QUÈ DIR
        say_heading = doc.add_paragraph()
        say_heading.paragraph_format.left_indent = Inches(0.2)
        say_run = say_heading.add_run("QUÈ DIR:")
        say_run.font.bold = True
        say_run.font.size = Pt(11)
        say_run.font.color.rgb = RGBColor(139, 90, 43)

        if slide.speaker_notes:
            notes_p = doc.add_paragraph()
            notes_p.paragraph_format.left_indent = Inches(0.4)
            _add_text_with_bold(notes_p, slide.speaker_notes, font_size=11)
        else:
            # Generar guia bàsica segons el tipus
            notes_p = doc.add_paragraph()
            notes_p.paragraph_format.left_indent = Inches(0.4)
            if slide.slide_type == "title":
                text = "Presenta't i introdueix el tema. Explica breument què veureu a la presentació."
            elif slide.slide_type == "index":
                text = "Explica l'estructura de la presentació i els punts principals que tractareu."
            else:
                text = "Explica els punts principals de la diapositiva amb les teves paraules."
            notes_p.add_run(text).font.size = Pt(11)

        # PUNTS A MENCIONAR
        if slide.content and slide.slide_type not in ["title"]:
            points_heading = doc.add_paragraph()
            points_heading.paragraph_format.left_indent = Inches(0.2)
            points_run = points_heading.add_run("PUNTS A MENCIONAR:")
            points_run.font.bold = True
            points_run.font.size = Pt(11)
            points_run.font.color.rgb = RGBColor(139, 90, 43)

            for item in slide.content:
                if not item or not item.strip():
                    continue
                text = item.strip()
                # Netejar markers de format
                text = text.replace("**", "")
                if text.isupper() and len(text) < 50:
                    # És un subtítol
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    run = p.add_run(f"→ {text}")
                    run.font.bold = True
                    run.font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.4)
                    run = p.add_run(f"• {text[:80]}{'...' if len(text) > 80 else ''}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(80, 80, 80)

        # Separador
        sep = doc.add_paragraph()
        sep.add_run("─" * 60).font.color.rgb = RGBColor(200, 200, 200)

    # ═══════════════════════════════════════════════════════════════════
    # CONSELLS FINALS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("CONSELLS PER LA PRESENTACIÓ", level=1)

    tips = [
        ("Abans de presentar", [
            "Llegeix els apunts complets almenys 2 vegades",
            "Practica en veu alta cronometrant-te",
            "Prepara possibles preguntes que et puguin fer"
        ]),
        ("Durant la presentació", [
            "Parla amb seguretat i mira al públic",
            "No llegeixis les diapositives, explica amb les teves paraules",
            "Cada slide ~1 minut, no et precipitis",
            "Connecta els conceptes entre ells"
        ]),
        ("Si et pregunten", [
            "Escolta tota la pregunta abans de respondre",
            "Si no saps la resposta, reconeix-ho i ofereix buscar-ho",
            "Relaciona la resposta amb el contingut de la presentació"
        ])
    ]

    for section_title, section_tips in tips:
        doc.add_heading(section_title, level=2)
        for tip in section_tips:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(tip).font.size = Pt(11)

    # Guardar
    doc.save(str(output_path))
    print(f"Apunts guardats: {output_path}")
    return output_path


def _add_text_with_bold(paragraph, text: str, font_size: int = 11):
    """
    Afegeix text a un paràgraf processant la sintaxi **negreta**.
    """
    import re

    if "**" not in text:
        run = paragraph.add_run(text)
        run.font.size = Pt(font_size)
        return

    # Processar text amb negreta
    pattern = r'\*\*([^*]+)\*\*'
    parts = re.split(pattern, text)

    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(font_size)
        if i % 2 == 1:
            run.font.bold = True


def _setup_styles(doc: Document):
    """Configura els estils del document."""
    # Estil normal
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Títol principal
    style = doc.styles['Heading 1']
    style.font.name = 'Calibri Light'
    style.font.size = Pt(18)
    style.font.color.rgb = RGBColor(224, 122, 47)
    style.font.bold = True

    # Subtítols
    style = doc.styles['Heading 2']
    style.font.name = 'Calibri'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(139, 90, 43)
    style.font.bold = True


def _add_compact_section(doc, section_title, content_list):
    """Afegeix una secció compacta amb bullet points."""
    if not content_list:
        return

    # Títol de secció
    section_p = doc.add_paragraph()
    section_run = section_p.add_run(section_title)
    section_run.font.bold = True
    section_run.font.size = Pt(12)
    section_run.font.color.rgb = RGBColor(224, 122, 47)
    section_p.paragraph_format.space_after = Pt(3)

    # Contingut com bullet points compactes
    for item in content_list[:8]:  # Màxim 8 punts per secció
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item).font.size = Pt(10)

    # Espai mínim
    doc.add_paragraph()
